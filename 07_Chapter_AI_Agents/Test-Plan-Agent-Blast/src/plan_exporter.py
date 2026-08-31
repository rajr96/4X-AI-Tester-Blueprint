"""Binary export helpers for generated test plans."""

from __future__ import annotations

import io

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from src.test_plan_generator import TestPlan


def to_docx(plan: TestPlan) -> bytes:
    document = Document()
    document.add_heading(f"Test Plan: {plan.issue_key}", level=0)
    document.add_paragraph(f"Summary: {plan.summary}")
    document.add_paragraph(f"Objective: {plan.objective}")
    _add_word_list(document, "Scope", plan.scope)
    _add_word_list(document, "Risks", plan.risks)
    document.add_heading("Test Cases", level=1)

    for case in plan.test_cases:
        document.add_heading(f"{case.test_id}: {case.title}", level=2)
        document.add_paragraph(f"Category: {case.category}")
        document.add_paragraph("Steps:")
        for step in case.steps:
            document.add_paragraph(step, style="List Bullet")
        document.add_paragraph(f"Expected result: {case.expected_result}")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def to_pdf(plan: TestPlan) -> bytes:
    output = io.BytesIO()
    styles = getSampleStyleSheet()
    story = [
        Paragraph(f"Test Plan: {plan.issue_key}", styles["Title"]),
        Paragraph(f"<b>Summary:</b> {plan.summary}", styles["BodyText"]),
        Paragraph(f"<b>Objective:</b> {plan.objective}", styles["BodyText"]),
        Spacer(1, 12),
    ]
    _add_pdf_list(story, "Scope", plan.scope, styles)
    _add_pdf_list(story, "Risks", plan.risks, styles)
    story.append(Paragraph("Test Cases", styles["Heading1"]))
    for case in plan.test_cases:
        story.append(Paragraph(f"{case.test_id}: {case.title}", styles["Heading2"]))
        story.append(Paragraph(f"<b>Category:</b> {case.category}", styles["BodyText"]))
        for step in case.steps:
            story.append(Paragraph(f"- {step}", styles["BodyText"]))
        story.append(Paragraph(f"<b>Expected result:</b> {case.expected_result}", styles["BodyText"]))
        story.append(Spacer(1, 8))

    SimpleDocTemplate(output, pagesize=letter).build(story)
    return output.getvalue()


def _add_word_list(document: Document, heading: str, items: list[str]) -> None:
    document.add_heading(heading, level=1)
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _add_pdf_list(story: list, heading: str, items: list[str], styles: dict) -> None:
    story.append(Paragraph(heading, styles["Heading1"]))
    for item in items:
        story.append(Paragraph(f"- {item}", styles["BodyText"]))
    story.append(Spacer(1, 8))